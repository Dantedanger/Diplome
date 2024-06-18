import sys

from PySide6.QtWidgets import QApplication, QFileDialog, QMessageBox, QMainWindow, QWidget, QLineEdit, QTextEdit, QTableWidget, QTableWidgetItem, QDialog
from PySide6.QtUiTools import QUiLoader
from PySide6.QtCore import QFile, QIODevice
from docx import Document
import mysql.connector as mc
import re

#     pyside6-uic D:\QtProjects\Diplom\form.ui -o D:\QtProjects\Diplom\ui_form.py
#     pyside6-uic D:\QtProjects\Diplom\discipline.ui -o D:\QtProjects\Diplom\ui_discipline.py
#     pyside6-uic D:\QtProjects\Diplom\competence.ui -o D:\QtProjects\Diplom\ui_competence.py
#     pyside6-uic D:\QtProjects\Diplom\edicational_standart.ui -o D:\QtProjects\Diplom\ui_edicational_standart.py
#     pyside6-uic D:\QtProjects\Diplom\desciplinecompetence.ui -o D:\QtProjects\Diplom\ui_desciplinecompetence.py
#     pyside6-uic D:\QtProjects\Diplom\desciplinecompetenceadd.ui -o D:\QtProjects\Diplom\ui_desciplinecompetenceadd.py
#     pyside6-uic D:\QtProjects\Diplom\educational_program.ui -o D:\QtProjects\Diplom\ui_educational_program.py
#     pyside6-uic D:\QtProjects\Diplom\educational_programadd.ui -o D:\QtProjects\Diplom\ui_educational_programadd.py
#     pyside6-uic D:\QtProjects\Diplom\syllibusdiscipline.ui -o D:\QtProjects\Diplom\ui_syllibusdiscipline.py
#     pyside6-uic D:\QtProjects\Diplom\syllibusdisciplineadd.ui -o D:\QtProjects\Diplom\ui_syllibusdisciplineadd.py
#     pyside6-uic D:\QtProjects\Diplom\result.ui -o D:\QtProjects\Diplom\ui_syllibusresult.py
#     pyside6-uic D:\QtProjects\Diplom\resultadd.ui -o D:\QtProjects\Diplom\ui_syllibusresulteadd.py

from ui_form import Ui_MainWindow
from discipline import Discipline
from competence import Competence
from edicational_standart import Edicational_standart
from desciplinecompetence import DesciplineCompetence
from educational_program import Educational_program
from syllibusdiscipline import SyllibusDiscipline
from result import Result

class MainWindow(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        self.ui.pushButton.clicked.connect(self.load_files)

        self.ui.action.triggered.connect(self.onActionTriggered)
        self.ui.action_2.triggered.connect(self.onAction_2Triggered)
        self.ui.action_3.triggered.connect(self.onAction_3Triggered)
        self.ui.action_4.triggered.connect(self.onAction_4Triggered)
        self.ui.action_5.triggered.connect(self.onAction_5Triggered)
        self.ui.action_6.triggered.connect(self.onAction_6Triggered)
        self.ui.action_7.triggered.connect(self.onAction_7Triggered)

        self.mydb = self.connect_database()

    def onActionTriggered(self):
        self.discipline_widget = Discipline(self.mydb)
        self.discipline_widget.setWindowTitle("Дисциплины")
        self.discipline_widget.exec()

    def onAction_2Triggered(self):
        self.discipline_widget = SyllibusDiscipline(self.mydb)
        self.discipline_widget.setWindowTitle("Дисциплины учебного плана")
        self.discipline_widget.exec()

    def onAction_3Triggered(self):
        self.competence_widget = Competence(self.mydb)
        self.competence_widget.setWindowTitle("Компетенции")
        self.competence_widget.exec()

    def onAction_4Triggered(self):
        self.competence_widget = DesciplineCompetence(self.mydb)
        self.competence_widget.setWindowTitle("Компетенции и дисциплины")
        self.competence_widget.exec()

    def onAction_5Triggered(self):
        self.edicational_standart_widget = Edicational_standart(self.mydb)
        self.edicational_standart_widget.setWindowTitle("Образовательный стандарт")
        self.edicational_standart_widget.exec()

    def onAction_6Triggered(self):
        self.edicational_standart_widget = Educational_program(self.mydb)
        self.edicational_standart_widget.setWindowTitle("Образовательная программа")
        self.edicational_standart_widget.exec()

    def onAction_7Triggered(self):
        self.edicational_standart_widget = Result(self.mydb)
        self.edicational_standart_widget.setWindowTitle("Связь двух дисциплин")
        self.edicational_standart_widget.exec()

    def connect_database(self):
        try:
            mydb = mc.connect(
                host = "localhost",
                user = "root",
                password = "",
                database = "09.03.04.database",
                charset='utf8mb4'
                )
#            mydb = sqlite3.connect('09_03_04_database_202404210022.sql')
            print(f'Connect')
            return mydb
        except mc.Error as e:
            print(f'NOT Connected')
            return False

    def closeEvent(self, event):
        self.mydb.close()
        print(f'Disconnect')
        event.accept()

    def addEdSt(self, prog, years):
        numb = prog[:8]
        prog = prog[9:]
        cursor = self.mydb.cursor()
        query = "INSERT INTO edicational_standart (Specialization_code, Name, Time) VALUES (%s, %s, %s)"
        try:
            cursor.execute(query,(numb, prog, years.lower()))
        except mc.Error as err:
            if err.errno == 1062:  # Ошибка нарушения уникального ключа
                pass  # Пропускаем эту ошибку
            else:
                print("Ошибка MySQL:", err.msg)
        self.mydb.commit()
        cursor.close()

    def addProfile(self, prog, profile, years):
        ided = 0
        cursor = self.mydb.cursor()
        query = "SELECT IDEdSt FROM edicational_standart WHERE Name = %s LIMIT 1"
        cursor.execute(query, (prog,))
        row = cursor.fetchone()  # Получение первой строки результата запроса
        if row:
            ided = row[0]
        cursor.close()

        cursor = self.mydb.cursor()
        query = "INSERT INTO edicational_program (IDEdSt, Profile, Year) VALUES (%s, %s, %s)"
        try:
            cursor.execute(query,(ided, profile, years))
        except mc.Error as err:
            if err.errno == 1062:  # Ошибка нарушения уникального ключа
                pass  # Пропускаем эту ошибку
            else:
                print("Ошибка MySQL:", err.msg)
        self.mydb.commit()
        cursor.close()

    def addDisciplineinProfile(self, profile, discipline, semester):
        idpr = 0
        idd = 0
        cursor = self.mydb.cursor()
        query = "SELECT IDEdPr FROM edicational_program WHERE Profile = %s LIMIT 1"
        cursor.execute(query, (profile,))
        row = cursor.fetchone()  # Получение первой строки результата запроса
        if row:
            idpr = row[0]
        cursor.close()

        cursor = self.mydb.cursor()
        query = "SELECT IDD FROM discipline WHERE Name = %s LIMIT 1"
        cursor.execute(query, (discipline,))
        row = cursor.fetchone()  # Получение первой строки результата запроса
        if row:
            idd = row[0]
        cursor.close()

        cursor = self.mydb.cursor()
        query = "INSERT INTO syllibusdiscipline (IDEdPr, IDD, Semesters) VALUES (%s, %s, %s)"
        try:
            cursor.execute(query,(idpr, idd, semester))
        except mc.Error as err:
            if err.errno == 1062:  # Ошибка нарушения уникального ключа
                pass  # Пропускаем эту ошибку
            else:
                print("Ошибка MySQL:", err.msg)
        self.mydb.commit()
        cursor.close()

    def addDiscipline(self, name, hour,descr):
        cursor = self.mydb.cursor()
        query = "INSERT INTO discipline (Name, QuantityAcademicHour, Description, Object) VALUES (%s, %s, %s, '')"
        try:
            cursor.execute(query,(name.lower(), hour, descr))
        except mc.Error as err:
            if err.errno == 1062:  # Ошибка нарушения уникального ключа
                pass  # Пропускаем эту ошибку
            else:
                print("Ошибка MySQL:", err.msg)
        self.mydb.commit()
        cursor.close()

    def addCompetence(self, name, desc):
        type = name[:name.find('-')]
        cursor = self.mydb.cursor()
        query = "INSERT INTO competence (Name, Description, Type) VALUES (%s, %s, %s)"
        try:
            cursor.execute(query,(name.upper(), desc, type))
        except mc.Error as err:
            if err.errno == 1062:  # Ошибка нарушения уникального ключа
                pass  # Пропускаем эту ошибку
            else:
                print("Ошибка MySQL:", err.msg)
        self.mydb.commit()
        cursor.close()

    def addCompetenceDiscipline(self, competence, discipline):
        idc = 0
        idd = 0
        cursor = self.mydb.cursor()
        query = "SELECT IDC FROM competence WHERE Name = %s LIMIT 1"
        cursor.execute(query, (competence,))
        row = cursor.fetchone()  # Получение первой строки результата запроса
        if row:
            idc = row[0]
        cursor.close()

        cursor = self.mydb.cursor()
        query = "SELECT IDD FROM discipline WHERE Name = %s LIMIT 1"
        cursor.execute(query, (discipline,))
        row = cursor.fetchone()  # Получение первой строки результата запроса
        if row:
            idd = row[0]
        cursor.close()

        cursor = self.mydb.cursor()
        query = "INSERT INTO desciplinecompetence (IDC, IDD) VALUES (%s, %s)"
        try:
            cursor.execute(query,(idc, idd))
        except mc.Error as err:
            if err.errno == 1062:  # Ошибка нарушения уникального ключа
                pass  # Пропускаем эту ошибку
            else:
                print("Ошибка MySQL:", err.msg)
        self.mydb.commit()
        cursor.close()

#    def firstly(self, doc_name):
#        doc = Document(doc_name)
#        doc_inner = doc._element
#        target_xpath = '//w:sdt//w:t'
#        for result_el in doc_inner.xpath(target_xpath):
#            result_el.text = ' '
#        doc.save(doc_name)

    def load_files(self):
        file_dialog = QFileDialog()
        file_dialog.setFileMode(QFileDialog.ExistingFiles)
        file_dialog.setNameFilter("Документы (*.docx)")

        if file_dialog.exec():
            files = file_dialog.selectedFiles()
            for file_name in files:
#                self.firstly(file_name)
                programm_discipline = self.search_words(file_name, "Направление подготовки")
                self.addEdSt(programm_discipline, self.search_words(file_name, "Уровень высшего образования"))

                discipline_name = self.search_words(file_name, "РАБОЧАЯ ПРОГРАММА ДИСЦИПЛИНЫ")
                self.addDiscipline(discipline_name, self.find_number_before_word(file_name, "час"),self.find_column_text_below_cell(file_name, "раздел/тема дисциплины"))

                for row_data in self.find_column_data_below_target(file_name, "Код компе-тенции"):
                    self.addCompetence(row_data[0], row_data[1])
                    self.addCompetenceDiscipline(row_data[0], discipline_name)

                programm_discipline = programm_discipline[9:]
                profile_discipline = self.search_words(file_name, "Направленность (профиль) образовательной программы")
                self.addProfile(programm_discipline, profile_discipline, self.search_numb(file_name,"решением Ученого совета"))

                self.addDisciplineinProfile(profile_discipline, discipline_name, self.find_column_below_target(file_name, "Семестр"))

                print(self.find_column_text_below_cell(file_name, "раздел/тема дисциплины"))
#print(self.search_words(file_name, "РАБОЧАЯ ПРОГРАММА ДИСЦИПЛИНЫ"))
#                print(self.search_words(file_name, "Направление подготовки"))
#                print(self.search_words(file_name, "Уровень высшего образования"))
#                print(self.search_words(file_name, "Направленность (профиль) образовательной программы"))

#                for row_data in self.find_column_data_below_target(file_name, "Код компе-тенции"):
#                    print(row_data[0], row_data[1], end = ' ')
#                print(self.find_number_before_word(file_name, "час"))

    def find_number_before_word(self, doc_name, target_word):
        doc = Document(doc_name)
        for paragraph in doc.paragraphs:
            match = re.search(r'(\d+)\s+' + re.escape(target_word), paragraph.text)
            if match:
                number = match.group(1)
                return number

    def find_column_below_target(self, doc_name, target_word):
        doc = Document(doc_name)
        for table in doc.tables:
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    if target_word in cell.text:
                        # Найдена ячейка с словом "Семестр"
                        # Проверяем, если это последний столбец таблицы
                        if cell_index == len(row.cells) - 1:
                            # Если это последний столбец, проверяем, есть ли следующая строка
                            if row_index < len(table.rows) - 1:
                                # Возвращаем текст ячейки ниже найденной ячейки
                                return table.cell(row_index + 1, cell_index).text.strip()
                        else:
                            # Возвращаем текст ячейки в том же столбце следующей строки
                            return table.cell(row_index + 1, cell_index).text.strip()
        return None


    def search_words(self, doc_name, target_word):
        doc = Document(doc_name)
        found_target = False

        for paragraph in doc.paragraphs:
            if found_target:
                if paragraph.text.strip():
                    return paragraph.text.strip()
            elif target_word in paragraph.text:
                found_target = True

    def search_numb(self, doc_name, start_phrase):
        doc = Document(doc_name)
        found_numbers = []
        end_phrase = "г."
        for paragraph in doc.paragraphs:
            # Ищем вхождения начальной и конечной фразы в тексте абзаца
            start_index = paragraph.text.find(start_phrase)
            end_index = paragraph.text.find(end_phrase)
            if start_index != -1 and end_index != -1:
                # Используем регулярное выражение для поиска чисел между фразами
                numbers_between_phrases = re.findall(r'\d+', paragraph.text[start_index:end_index])
                # Если нашли числа, добавляем их в список найденных чисел
                if numbers_between_phrases:
                    found_numbers.extend(map(int, numbers_between_phrases))

        # Если были найдены числа между фразами, возвращаем последнее из них
        if found_numbers:
            return max(found_numbers)
        else:
            return None


    def find_word_after_target(self, doc_name, target_word):
        doc = Document(doc_name)
        found_target = False

        for paragraph in doc.paragraphs:
            if found_target:
                words = paragraph.text.strip().split()
                if words:
                    next_word = words[0]
                    return next_word
            elif target_word in paragraph.text:
                found_target = True

    def find_column_data_below_target(self, doc_name, target_text):
        doc = Document(doc_name)
        found_target = False
        target_row_data = []

        for table in doc.tables:
            for row in table.rows:
                if found_target:
                    first_column_data = row.cells[0].text.strip()
                    second_column_data = row.cells[1].text.strip()
                    target_row_data.append((first_column_data, second_column_data))
                else:
                    for cell in row.cells:
                        if target_text in cell.text.strip():
                            found_target = True
                            break
            if found_target:
                found_target = False
                break
        return target_row_data

    def find_column_text_below_cell(self, doc_name, cell_text):
        column_text = []
        doc = Document(doc_name)
        for table in doc.tables:
            for row in table.rows:

                for cell_index, cell in enumerate(row.cells):
                    # Проверяем, если текст ячейки соответствует искомому
                    if cell_text in cell.text:
                        # Определяем индекс столбца, в котором находится искомая ячейка
                        column_index = cell_index

                        # Собираем текст из всех ячеек этого столбца, начиная со следующей строки
                        for i in range(row._index + 1, len(table.rows)):
                            b = table.rows[i].cells[column_index].text.strip()
                            if not(b in column_text):
                                column_text.append(b)
                        return str(' '.join(column_text))


if __name__ == "__main__":
    app = QApplication(sys.argv)
    widget = MainWindow()
    widget.show()
    sys.exit(app.exec())
